from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt


INTRO_BLOCKS = [
    ("heading", "Lý do chọn đề tài"),
    (
        "body",
        "Trong bối cảnh chuyển đổi số giáo dục đang diễn ra mạnh mẽ, việc học tập của sinh viên ngày càng gắn liền với các nền tảng số, tài liệu điện tử và các công cụ hỗ trợ học tập trực tuyến. Bên cạnh giáo trình truyền thống, sinh viên thường phải tiếp cận nhiều nguồn tài liệu khác nhau như file PDF, tài liệu DOCX, ghi chú bài giảng, slide môn học và các tài liệu tham khảo được cung cấp trong quá trình học. Khối lượng tài liệu lớn giúp người học có nhiều nguồn tri thức hơn, nhưng đồng thời cũng tạo ra khó khăn trong việc tìm kiếm nhanh nội dung cần thiết, tổng hợp kiến thức và tra cứu lại thông tin theo từng môn học.",
    ),
    (
        "body",
        "Trong những năm gần đây, các mô hình ngôn ngữ lớn đã cho thấy khả năng hỗ trợ người dùng đặt câu hỏi và nhận câu trả lời bằng ngôn ngữ tự nhiên. Tuy nhiên, nếu chỉ sử dụng mô hình ngôn ngữ một cách độc lập, câu trả lời có thể không bám sát tài liệu học tập cụ thể của sinh viên, thậm chí có nguy cơ tạo ra thông tin không chính xác. Điều này đặc biệt quan trọng trong môi trường học tập, nơi câu trả lời cần dựa trên nội dung tài liệu chính thức, có nguồn trích dẫn rõ ràng và phù hợp với ngữ cảnh môn học.",
    ),
    (
        "body",
        "Retrieval-Augmented Generation (RAG) là một hướng tiếp cận phù hợp để giải quyết vấn đề trên. Thay vì để mô hình ngôn ngữ trả lời hoàn toàn dựa trên tri thức tổng quát, hệ thống RAG thực hiện tìm kiếm các đoạn nội dung liên quan trong tài liệu đã được xử lý, sau đó sử dụng các đoạn nội dung này làm ngữ cảnh để sinh câu trả lời. Cách tiếp cận này giúp câu trả lời bám sát tài liệu hơn, đồng thời có thể cung cấp nguồn tham chiếu như tên tài liệu và số trang. Vì vậy, đề tài \"Xây dựng hệ thống chatbot hỏi đáp tài liệu học tập có hỗ trợ cá nhân hóa cho sinh viên sử dụng Retrieval-Augmented Generation\" được lựa chọn nhằm nghiên cứu và xây dựng một hệ thống hỗ trợ sinh viên tra cứu kiến thức học tập hiệu quả hơn.",
    ),
    ("heading", "Mục tiêu của đề tài"),
    (
        "body",
        "Mục tiêu tổng quát của đề tài là xây dựng một hệ thống chatbot hỗ trợ sinh viên hỏi đáp dựa trên tài liệu học tập đã được cung cấp. Hệ thống cho phép người dùng tải lên tài liệu, xử lý nội dung tài liệu, tạo dữ liệu vector phục vụ tìm kiếm ngữ nghĩa, truy xuất các đoạn nội dung liên quan đến câu hỏi và sử dụng mô hình ngôn ngữ để tạo câu trả lời có kèm nguồn tham khảo.",
    ),
    ("body", "Các mục tiêu cụ thể của đề tài gồm:"),
    ("bullet", "Xây dựng backend API bằng FastAPI để quản lý các chức năng chính của hệ thống."),
    ("bullet", "Thiết kế cơ sở dữ liệu lưu thông tin người dùng, khóa học, tài liệu và lịch sử hỏi đáp."),
    ("bullet", "Xây dựng chức năng tải lên tài liệu học tập ở các định dạng phổ biến như PDF, DOCX và TXT."),
    ("bullet", "Xây dựng module đọc, làm sạch và chuẩn hóa nội dung tài liệu sau khi tải lên."),
    ("bullet", "Chia nội dung tài liệu thành các đoạn nhỏ phù hợp cho quá trình tìm kiếm ngữ nghĩa."),
    ("bullet", "Tạo embedding cho các đoạn văn bản bằng mô hình hỗ trợ đa ngôn ngữ, phù hợp với tài liệu tiếng Việt."),
    ("bullet", "Lưu trữ embedding và metadata tài liệu trong ChromaDB để phục vụ truy xuất."),
    ("bullet", "Xây dựng chức năng tìm kiếm các đoạn tài liệu liên quan theo câu hỏi của người dùng."),
    ("bullet", "Tích hợp mô hình ngôn ngữ Gemini để sinh câu trả lời dựa trên ngữ cảnh được truy xuất."),
    ("bullet", "Trả lời câu hỏi kèm nguồn tham khảo và lưu lại lịch sử hỏi đáp để làm nền tảng cho chức năng cá nhân hóa về sau."),
    ("heading", "Phương pháp thực hiện"),
    (
        "body",
        "Đề tài được thực hiện theo hướng kết hợp giữa nghiên cứu lý thuyết và xây dựng thử nghiệm hệ thống. Trước hết, đề tài tìm hiểu các kiến thức nền tảng liên quan đến chatbot, mô hình ngôn ngữ lớn, embedding, vector database và kiến trúc Retrieval-Augmented Generation. Trên cơ sở đó, hệ thống được phân tích thành các module nhỏ như xử lý tài liệu, chia đoạn văn bản, tạo embedding, lưu trữ vector, truy xuất dữ liệu và sinh câu trả lời.",
    ),
    (
        "body",
        "Trong quá trình triển khai, đề tài áp dụng phương pháp thực nghiệm theo từng giai đoạn. Mỗi chức năng được xây dựng và kiểm thử độc lập trước khi tích hợp vào pipeline hoàn chỉnh. Các endpoint API được kiểm tra thông qua Swagger UI; dữ liệu tài liệu sau xử lý được kiểm tra thủ công nhằm đánh giá khả năng đọc tiếng Việt, giữ số trang và loại bỏ các phần nội dung không cần thiết. Đối với chức năng truy xuất, hệ thống được kiểm tra bằng các câu hỏi có trong tài liệu để đánh giá mức độ liên quan của các đoạn văn bản được trả về. Đối với chức năng chatbot, hệ thống được kiểm tra bằng cả câu hỏi có trong tài liệu và câu hỏi ngoài phạm vi tài liệu nhằm đánh giá khả năng trả lời dựa trên nguồn tham khảo.",
    ),
    (
        "body",
        "Ngoài ra, đề tài sử dụng Git và GitHub để quản lý mã nguồn theo từng giai đoạn phát triển. Việc chia nhỏ công việc theo tuần giúp quá trình thực hiện dễ kiểm soát hơn, đồng thời tạo minh chứng rõ ràng cho từng kết quả đã hoàn thành.",
    ),
    ("heading", "Đối tượng và phạm vi thực hiện"),
    (
        "body",
        "Đối tượng thực hiện của đề tài là hệ thống chatbot hỗ trợ hỏi đáp tài liệu học tập cho sinh viên. Trọng tâm của hệ thống là quy trình xử lý tài liệu và truy xuất thông tin phục vụ trả lời câu hỏi theo kiến trúc RAG. Các thành phần chính bao gồm quản lý người dùng, quản lý khóa học, tải lên tài liệu, xử lý văn bản, tạo embedding, lưu vector, tìm kiếm ngữ nghĩa, tích hợp mô hình ngôn ngữ và lưu lịch sử hỏi đáp.",
    ),
    (
        "body",
        "Phạm vi thực hiện của đề tài tập trung vào xây dựng phiên bản thử nghiệm phục vụ minh họa chức năng cốt lõi. Hệ thống hỗ trợ các định dạng tài liệu phổ biến gồm PDF, DOCX và TXT; sử dụng SQLite cho cơ sở dữ liệu ở giai đoạn đầu; sử dụng ChromaDB làm vector database; sử dụng mô hình sentence-transformers đa ngôn ngữ để tạo embedding; và sử dụng Gemini API để sinh câu trả lời. Chức năng cá nhân hóa được định hướng thông qua việc lưu lịch sử hỏi đáp, thông tin người dùng và khóa học, tuy nhiên chưa đi sâu vào các thuật toán cá nhân hóa nâng cao như phân tích năng lực học tập, gợi ý lộ trình học hoặc điều chỉnh câu trả lời theo hồ sơ học tập chi tiết.",
    ),
    (
        "body",
        "Đề tài không tập trung xây dựng một hệ quản trị học tập hoàn chỉnh, không xử lý toàn bộ các nghiệp vụ phức tạp của LMS và chưa đánh giá trên tập dữ liệu lớn ở quy mô triển khai thực tế. Thay vào đó, đề tài ưu tiên xây dựng pipeline RAG hoàn chỉnh, có khả năng chạy thử, kiểm chứng được kết quả và có thể mở rộng trong các giai đoạn tiếp theo.",
    ),
    ("heading", "Nội dung của báo cáo"),
    (
        "body",
        "Báo cáo được tổ chức thành phần mở đầu, các chương nội dung chính và phần kết luận. Cụ thể:",
    ),
    ("bullet", "Phần mở đầu trình bày lý do chọn đề tài, mục tiêu, phương pháp thực hiện, đối tượng và phạm vi thực hiện, cùng bố cục tổng quát của báo cáo."),
    ("bullet", "Chương 1 trình bày tổng quan về đề tài và cơ sở lý thuyết liên quan, bao gồm chatbot, mô hình ngôn ngữ lớn, embedding, vector database và kiến trúc Retrieval-Augmented Generation."),
    ("bullet", "Chương 2 trình bày phân tích yêu cầu và thiết kế hệ thống, bao gồm yêu cầu chức năng, yêu cầu phi chức năng, kiến trúc tổng thể, thiết kế cơ sở dữ liệu và thiết kế các thành phần xử lý chính."),
    ("bullet", "Chương 3 trình bày quá trình triển khai hệ thống, bao gồm xây dựng backend API, xử lý tài liệu, tạo chunk, embedding, lưu trữ vector, retrieval, tích hợp LLM và xây dựng API chat."),
    ("bullet", "Chương 4 trình bày kết quả thực nghiệm, minh chứng chức năng, đánh giá kết quả đạt được, các hạn chế còn tồn tại và hướng cải thiện."),
    ("bullet", "Phần kết luận tổng hợp các kết quả đã đạt được, rút ra bài học trong quá trình thực tập và đề xuất hướng phát triển tiếp theo cho hệ thống."),
]


def insert_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def set_common_format(paragraph: Paragraph, *, bold: bool = False, bullet: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(3)
    if bullet:
        fmt.left_indent = Pt(18)
        fmt.first_line_indent = Pt(-9)
    for run in paragraph.runs:
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: update_report_intro.py <input.docx> <output.docx>")

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    doc = Document(input_path)

    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "LỜI MỞ ĐẦU")
    end = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("TỔNG QUAN ĐỀ TÀI"))

    for para in paragraphs[start + 1 : end]:
        para._element.getparent().remove(para._element)

    anchor = doc.paragraphs[start]
    last = anchor
    for kind, text in INTRO_BLOCKS:
        if kind == "heading":
            para = insert_after(last, text)
            set_common_format(para, bold=True)
        elif kind == "bullet":
            para = insert_after(last, f"- {text}")
            set_common_format(para, bullet=True)
        else:
            para = insert_after(last, text)
            set_common_format(para)
        last = para

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    main()
